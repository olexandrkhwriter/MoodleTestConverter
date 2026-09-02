# -*- coding: utf-8 -*-
"""Інтеграційний тест: сирі файли різних форматів -> .mbz через mbz_module."""
import gzip
import io
import os
import sys
import tarfile
import xml.etree.ElementTree as ET

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from course_module import GlobalSettings
from mbz_module import build_mbz_from_files, files_to_course

# --- сирі файли різних форматів ------------------------------------------
os.makedirs("/tmp/mbz_mixed", exist_ok=True)

# 1) звичайний txt (Aiken)
with open("/tmp/mbz_mixed/01_слр.txt", "w", encoding="utf-8") as f:
    f.write("Співвідношення компресій до вдихів при СЛР дорослого?\n"
            "A. 15:2\nB. 30:2\nC. 5:1\nD. 10:1\nANSWER: B\n\n"
            "Перша дія при клінічній смерті?\n"
            "A. Виклик швидкої\nB. Початок СЛР\n"
            "C. Визначення пульсу\nD. Введення адреналіну\nANSWER: B\n")

# 2) docx зі зірочками (через converter_core -> to_aiken)
from docx import Document
doc = Document()
doc.add_paragraph("newTem;")
doc.add_paragraph("name:М1 Тестова тема")
doc.add_paragraph("QuestName:Що є першою ланкою ланцюга виживання?")
doc.add_paragraph("дефібриляція")
doc.add_paragraph("*раннє розпізнавання та виклик допомоги")
doc.add_paragraph("СЛР")
doc.add_paragraph("інтенсивна терапія")
doc.add_paragraph("trueNum:2")
doc.save("/tmp/mbz_mixed/02_ланцюг.docx")

# 3) csv з однією колонкою (Aiken-рядки)
with open("/tmp/mbz_mixed/03_опіки.txt", "w", encoding="cp1251") as f:
    f.write("Дія при термічному опіку I ступеня?\n"
            "A. Змастити олією\nB. Охолодити проточною водою\n"
            "C. Проколоти пухирі\nD. Накласти вату\nANSWER: B\n")

st = GlobalSettings(course_name="Тактична медицина",
                    period_prefix="Заняття", time_limit=20, attempts=2,
                    grading_method="highest", pass_percent=70,
                    shuffle_answers=True)

paths = ["/tmp/mbz_mixed/01_слр.txt",
         "/tmp/mbz_mixed/02_ланцюг.docx",
         "/tmp/mbz_mixed/03_опіки.txt"]

course, report = files_to_course(paths, st,
                                 course_fullname="М1 Тактична медицина")
print(report)
assert len(course.sections) == 3, f"секцій: {len(course.sections)}"
# перевіряємо, що docx-питання розпізнане з правильною відповіддю
docx_quiz = course.sections[1].quizzes[0]
assert docx_quiz.questions, "docx: немає питань"
cq = docx_quiz.questions[0]
assert cq.correct_index >= 0, "docx: правильна відповідь не визначена"

out = "/tmp/mbz_mixed/курс_м1.mbz"
rep = build_mbz_from_files(paths, st, out,
                           course_fullname="М1 Тактична медицина",
                           review_preset="strict")
assert os.path.getsize(out) > 4000

# --- перевірка strict review flags ---
z = tarfile.open(fileobj=io.BytesIO(
    gzip.decompress(open(out, "rb").read())))
quiz1 = z.extractfile("activities/quiz_1/quiz.xml").read().decode("utf-8")
assert "<reviewcorrectness>0</reviewcorrectness>" in quiz1
assert "<reviewrightanswer>0</reviewrightanswer>" in quiz1
assert "<reviewmarks>4352</reviewmarks>" in quiz1
print("\nstrict review flags ✔")

# кириличні назви файлів усередині .mbz відсутні (лише ASCII-шляхи)
assert all(all(ord(c) < 128 for c in n) for n in z.getnames()), \
    "шляхи в .mbz мають бути ASCII"
print("ASCII-шляхи в архіві ✔")

# усі XML валідні
for n in z.getnames():
    if n.endswith(".xml"):
        ET.fromstring(z.extractfile(n).read().decode("utf-8"))
print("усі XML валідні ✔")

print(f"\n=== ІНТЕГРАЦІЙНИЙ ТЕСТ ПРОЙДЕНО: {out} "
      f"({os.path.getsize(out)/1024:.1f} КБ) ===")
