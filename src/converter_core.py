# -*- coding: utf-8 -*-
"""
Moodle Test Converter — core logic
Parse test files (.txt, .docx, .xlsx, .csv, .html) with marked correct answers
and export to Moodle formats: GIFT, Moodle XML, Aiken.
"""

import re
import csv
import html
import io
import os
import xml.sax.saxutils as saxutils
from dataclasses import dataclass, field
from typing import List, Optional

# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------

@dataclass
class Answer:
    text: str
    correct: bool = False
    fraction: float = 0.0          # 0..100 (Moodle percent weight)
    feedback: str = ""
    match: str = ""                # right side for matching questions


@dataclass
class Question:
    qtype: str = "multichoice"     # multichoice | truefalse | shortanswer | matching | numerical | essay | description
    name: str = ""
    text: str = ""
    answers: List[Answer] = field(default_factory=list)
    feedback: str = ""             # general feedback
    default_mark: float = 1.0


# ---------------------------------------------------------------------------
# File reading (encoding detection)
# ---------------------------------------------------------------------------

_ENCODINGS = ["utf-8-sig", "utf-8", "cp1251", "cp1252", "latin-1"]


def decode_bytes(data: bytes) -> str:
    for enc in _ENCODINGS:
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def read_text_file(path: str) -> str:
    with open(path, "rb") as f:
        return decode_bytes(f.read())


def read_docx_paragraphs(path: str):
    """Return (paragraphs, tables).
    paragraphs: list of (text, style_set) where style_set may contain
    'bold', 'hl' (highlighted), 'und' (underlined), 'color' (red font).
    tables: list of rows (list of cell texts) for Word table templates.
    """
    from docx import Document
    from docx.enum.text import WD_COLOR_INDEX

    doc = Document(path)
    out = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt:
            out.append(("", set()))
            continue
        marks = set()
        for r in p.runs:
            if r.text.strip():
                if r.bold:
                    marks.add("bold")
                if r.underline:
                    marks.add("und")
                try:
                    if r.font.highlight_color is not None and \
                       r.font.highlight_color != WD_COLOR_INDEX.AUTO:
                        marks.add("hl")
                except Exception:
                    pass
                try:
                    c = r.font.color
                    if c is not None and c.rgb is not None:
                        rgb = c.rgb
                        # red-ish font often marks the correct answer
                        if rgb[0] > 0x99 and rgb[1] < 0x66 and rgb[2] < 0x66:
                            marks.add("color")
                except Exception:
                    pass
        out.append((txt, marks))
    # tables in docx (university templates often use tables)
    tables = []
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                tables.append(cells)
    # Split multi-line paragraphs: a single Word paragraph may contain the
    # question AND all answer options separated by line breaks (\n / <w:br/>).
    # The star marker is often a separate run '*', glued to the option letter
    # ('*' + 'В. ...' -> '*В.'), so keep it attached when splitting.
    split_out = []
    for txt, marks in out:
        if "\n" not in txt:
            split_out.append((txt, marks))
            continue
        for part in txt.split("\n"):
            part = part.strip()
            if part:
                split_out.append((part, marks))
    return split_out, tables


def read_xlsx_rows(path: str):
    from openpyxl import load_workbook
    wb = load_workbook(path, read_only=True, data_only=True)
    rows = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            vals = [str(c).strip() if c is not None else "" for c in row]
            if any(vals):
                rows.append(vals)
    return rows


def read_csv_rows(path: str):
    text = read_text_file(path)
    # sniff delimiter
    try:
        dialect = csv.Sniffer().sniff(text[:2048], delimiters=";,\t")
    except Exception:
        dialect = csv.excel
    return [r for r in csv.reader(io.StringIO(text), dialect) if any(x.strip() for x in r)]


def _text_looks_broken(text: str) -> bool:
    """Detect a garbled extraction (encoding lost): the text is mostly '?'
    characters where Cyrillic/Latin letters should be. Typical symptom of a
    Windows code-page mismatch (e.g. LibreOffice txt export into a
    non-Cyrillic system locale turns every letter into '?')."""
    if not text or len(text) < 40:
        return False
    letters = sum(1 for c in text if c.isalpha())
    qmarks = text.count("?")
    # healthy files: '?' is rare (<2% of letters); garbled: '?' dominates
    if letters < 20 and qmarks > 20:
        return True
    return letters > 0 and qmarks > letters * 0.5


def read_doc_text(path: str) -> str:
    """Extract text from legacy binary .doc (Word 97-2003).
    Strategy order (whatever is available on the user's machine):
      1) LibreOffice headless conversion to .txt FORCED to UTF-8
         ('txt:Text (encoded):UTF8') — without the encoding token LO on
         Windows writes the system code page and Cyrillic becomes '?????'
      2) antiword with an explicit UTF-8 mapping file
      3) MS Word COM automation (if Word is installed)
      4) crude fallback: decode the WordDocument stream as UTF-16LE /
         cp1251 and keep printable runs (loses bold/highlight info, but
         star markers and the letter layout survive).
    Every intermediate result is validated with _text_looks_broken() so a
    garbled conversion never silently wins over a working fallback.
    """
    import shutil
    import subprocess
    import tempfile

    # 1) LibreOffice — always force UTF-8 output
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        for cand in (r"C:\Program Files\LibreOffice\program\soffice.exe",
                     r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"):
            if os.path.exists(cand):
                soffice = cand
                break
    if soffice:
        try:
            with tempfile.TemporaryDirectory() as tmp:
                subprocess.run(
                    [soffice, "--headless", "--convert-to",
                     "txt:Text (encoded):UTF8",
                     "--outdir", tmp, path],
                    capture_output=True, timeout=120)
                base = os.path.splitext(os.path.basename(path))[0]
                out = os.path.join(tmp, base + ".txt")
                if os.path.exists(out):
                    with open(out, "rb") as f:
                        text = decode_bytes(f.read())
                    if not _text_looks_broken(text):
                        return text
        except Exception:
            pass

    # 2) antiword (explicit UTF-8 mapping; falls back to default mapping)
    antiword = shutil.which("antiword")
    if antiword:
        for cmd in ([antiword, "-m", "UTF-8.txt", path],
                    [antiword, path]):
            try:
                out = subprocess.run(cmd, capture_output=True, timeout=60)
                if out.returncode == 0 and out.stdout:
                    text = decode_bytes(out.stdout)
                    if not _text_looks_broken(text):
                        return text
            except Exception:
                continue

    # 3) MS Word COM (Windows with Word installed)
    try:
        import win32com.client  # type: ignore
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        try:
            doc = word.Documents.Open(os.path.abspath(path), ReadOnly=True)
            text = doc.Content.Text
            doc.Close(False)
        finally:
            word.Quit()
        text = text.replace("\r", "\n")
        if not _text_looks_broken(text):
            return text
    except Exception:
        pass

    # 4) crude fallback — extract printable text from the binary stream
    with open(path, "rb") as f:
        data = f.read()
    # Word stores text mostly as UTF-16LE in the WordDocument stream
    best = ""
    for enc in ("utf-16-le", "cp1251", "utf-8"):
        try:
            raw = data.decode(enc, errors="ignore")
        except Exception:
            continue
        # keep runs of letters/digits/punctuation (drop binary noise)
        runs = re.findall(
            r'[\w\sА-Яа-яЇїІіЄєҐґ.,;:!?()\[\]\-–—«»„“”\'’/*+=%№<>→°]{6,}',
            raw, flags=re.UNICODE)
        text = "\n".join(r.strip() for r in runs if r.strip())
        # prefer the decoding that yields more Cyrillic
        if text.count("і") + text.count("а") > best.count("і") + best.count("а"):
            best = text
    return best


def read_html_text(path: str) -> str:
    text = read_text_file(path)
    text = re.sub(r"(?is)<(script|style).*?</\1>", "", text)
    text = re.sub(r"(?i)<br\s*/?>", "\n", text)
    text = re.sub(r"(?i)</(p|div|li|tr|h[1-6])>", "\n", text)
    text = re.sub(r"<[^>]+>", "", text)
    return html.unescape(text)


# ---------------------------------------------------------------------------
# Structure detection & parsing
# ---------------------------------------------------------------------------

# answer option: letter/digit + separator, OPTIONAL star glued to the
# letter ('*В. відповідь' — the star is a separate run in Word, no space)
_ANSWER_LETTER_RE = re.compile(
    r"^\s*([\*\+✓✔√]?)\s*([A-Za-zА-Яа-яЇїІіЄєҐґ]|\d{1,2})\s*[\.\)\:]\s+(.*)$")
# a plain unnumbered question line (medical dumps: no '1.' at all).
# A question is a line that CONTAINS '?' (it may sit mid-line when the
# question has several sentences) and is long enough to be a real question.
_PLAIN_QUESTION_RE = re.compile(r"[?？]")
_MARKER_PREFIX_RE = re.compile(r"^\s*([\+\-])\s+([A-Za-zА-Яа-яЇїІіЄєҐґ]|\d{1,2})\s*[\.\)\:\-–—]\s+(.*)$")
_QUESTION_NUM_RE = re.compile(
    r"^\s*(?:запитання|питання|question|завдання|q)?\s*[№#]?\s*(\d{1,4})\s*[\.\)\:\-–—]\s+(.*[^*+\s])\s*$",
    re.IGNORECASE)
_QUESTNAME_RE = re.compile(
    r"^\s*QuestName\s*[:：]\s*(.*)$", re.IGNORECASE)
_ANSWER_KEY_RE = re.compile(
    r"^\s*(?:ans(?:we|ve|wv|ev)?r|відповідь|відповіді|"
    r"правильн\w*(?:\s+відповід\w*)?|ключ\s+відповід\w*)\s*[:：\-–]?\s*(\S.*)?$",
    re.IGNORECASE)
_LETTER_KEY_RE = re.compile(
    r"^[A-Za-zА-Яа-яЇїІіЄєҐґ](\s*[,;\s]\s*[A-Za-zА-Яа-яЇїІіЄєҐґ])*\.?$",
    re.IGNORECASE)
_NUM_KEY_RE = re.compile(r"^\d{1,2}(\s*[,;\s]\s*\d{1,2})*\.?$")
_TAIL_KEY_LINE_RE = re.compile(
    r"^\s*(?:відповіді|answers|ключ)\s*[:：]\s*"
    r"(?:(?:\d+\s*[-–:=]\s*[A-Za-zА-Яа-яЇїІіЄєҐґ0-9]{1,2})[,;\s]*)+$",
    re.IGNORECASE)
_TRUEFALSE_RE = re.compile(
    r"\b(true|false|так|ні|вірно|невірно|правда|брехня)\b", re.IGNORECASE)
_TF_UA_WORDS = {"так": True, "ні": False, "вірно": True, "невірно": False,
                "правда": True, "брехня": False}


def _clean_correct_mark(text: str):
    """Remove leading/trailing correct-answer markers; return (text, was_marked)."""
    t = text.strip()
    marked = False
    # markdown bold first (before trailing-star strip): **answer** / __answer__
    m = re.fullmatch(r"\*\*(.+)\*\*|__(.+)__", t)
    if m:
        t = (m.group(1) or m.group(2)).strip()
        marked = True
    m = re.match(r"^[\+✓✔√☑]\s*", t)
    if m:
        t = t[m.end():].strip()
        marked = True
    # leading star that is NOT markdown (no closing ** later)
    if t.startswith("*") and "**" not in t:
        t = t[1:].strip()
        marked = True
    m = re.search(r"\s*[\*\+✓✔√]\s*$", t)
    if m and "**" not in t:
        t = t[:m.start()].strip()
        marked = True
    # trailing correctness note (with optional markdown stars around it):
    # 'text (правильна відповідь)' / 'text **(правильна відповідь)**'
    # (word boundary: 'Невірно' must NOT match 'вірно')
    # (English words require a word boundary so '15:2' does NOT match
    #  'correct'→'rect' inside digits, and 'CO2' etc. stay intact)
    m = re.search(
        r"\s*\*{0,2}[\(\[]\s*(вірно|правильно|правильна(?:\s+відповідь)?|correct)\s*[\)\]]\s*\*{0,2}$"
        r"|\s*\*{0,2}(правильно|правильна(?:\s+відповідь)?)\s*\*{0,2}$"
        r"|\s+\b(correct)\b\s*$",
        t, re.IGNORECASE)
    if m and len(t[:m.start()].strip()) > 0:
        t = t[:m.start()].strip()
        marked = True
    # strip leftover markdown emphasis chars
    t = re.sub(r"\*\*|__", "", t).strip()
    return t, marked


def _split_match(text: str):
    """Split 'left -> right' / 'left = right' matching pair."""
    for sep in ("->", "=>", "→", "—", "–", "="):
        if sep in text:
            l, r = text.split(sep, 1)
            if l.strip() and r.strip():
                return l.strip(), r.strip()
    return text.strip(), ""


def parse_numbered_lines(lines, style_marks=None):
    """
    Parse numbered-list layout:
        1. Question text
        a) answer
        b) answer *          (marked) / bold / highlighted
        ...
    Optional answer key at the end: 'Відповіді: 1-a, 2-b' or per-question
    'ANSWER: A' / 'Відповідь: Б'.
    style_marks: parallel list of sets with 'bold'/'hl'/'und' flags per line.
    """
    questions = []
    cur = None
    pending_key = {}  # qnum -> letter(s)
    tail_key = {}     # answers listed at end of file

    lines = list(lines)
    style_marks = style_marks or [set() for _ in lines]

    # --- pre-scan for an "answers at end" block like: Відповіді: 1-а; 2-б ...
    # also 'Ключ: 1-а; 2-б' or numeric 'Ключ: 1-2; 2-1'
    joined_tail = "\n".join(lines[-10:])
    km = re.search(
        r"(?:відповіді|answers|ключ)\s*[:：]\s*((?:\d+\s*[-–:=]\s*[A-Za-zА-Яа-яЇїІіЄєҐґ0-9]{1,2}\s*[,;\s]*)+)",
        joined_tail, re.IGNORECASE)
    if km:
        for qn, let in re.findall(
                r"(\d+)\s*[-–:=]\s*([A-Za-zА-Яа-яЇїІіЄєҐґ0-9]{1,2})", km.group(1)):
            tail_key[int(qn)] = let.lower()

    def flush():
        nonlocal cur
        if cur is not None and (cur.answers or cur.qtype in
                                ("essay", "description")):
            questions.append(cur)
        cur = None

    last_qnum = 0
    pending_title = None  # unnumbered line(s) right before question 1 = file title
    questname_mode = False   # 'QuestName:' blocks without numbering
    topic_title = None       # 'name:...' header of newTem; dumps
    for idx, raw in enumerate(lines):
        line = raw.rstrip()
        if not line.strip():
            continue
        marks = style_marks[idx]

        # skip a trailing 'Відповіді: 1-а, 2-б' key block
        if _TAIL_KEY_LINE_RE.match(line):
            continue

        # --- topic-header lines of the medical/exam dump format:
        #     newTem;
        #     name:Назва теми (becomes the Moodle category)
        if re.match(r"^\s*newTem\s*;?\s*$", line, re.IGNORECASE):
            continue
        nm = re.match(r"^\s*name\s*[:：]\s*(.+)$", line, re.IGNORECASE)
        if nm and cur is None and not questions:
            topic_title = nm.group(1).strip()
            continue

        # --- 'QuestName: <text>' block format (medical/exam dumps):
        #     QuestName: текст питання
        #     варіант 1
        #     варіант 2
        #     trueNum:3          (1-based index of the correct option)
        qnm = _QUESTNAME_RE.match(line)
        if qnm:
            questname_mode = True
            flush()
            cur = Question(name="", text=qnm.group(1).strip())
            cur._qnum = len(questions) + 1
            last_qnum = cur._qnum
            continue
        tnm = re.match(r"^\s*trueNum\s*[:：]\s*(\d{1,2})\s*$", line,
                       re.IGNORECASE)
        if tnm:
            if cur is not None:
                i = int(tnm.group(1)) - 1
                if 0 <= i < len(cur.answers):
                    cur.answers[i].correct = True
                    cur.answers[i].fraction = 100
            continue

        if questname_mode:
            # plain (unlettered) lines are answer options of the current
            # QuestName question; an empty question means end of the block.
            if cur is not None and cur.text and line.strip():
                # but a key line ('Ansver:2', 'Правильна відповідь: 2',
                # 'Відповідь: Б' etc.) marks the correct option by number
                # or letter and is NOT added as an answer option
                km2 = _ANSWER_KEY_RE.match(line.strip())
                if km2:
                    val = (km2.group(1) or "").strip().rstrip(".")
                    if val.isdigit():
                        i = int(val) - 1
                        if 0 <= i < len(cur.answers):
                            cur.answers[i].correct = True
                            cur.answers[i].fraction = 100
                    elif _LETTER_KEY_RE.match(val):
                        for j, a in enumerate(cur.answers):
                            a._letter = chr(ord("a") + j)
                        _apply_letters(cur, [val.lower()])
                    continue
                atext, marked = _clean_correct_mark(line.strip())
                cur.answers.append(
                    Answer(atext, marked, 100 if marked else 0))
            continue

        # ANSWER: X  /  Відповідь: Б   (per-question key)
        mk = _ANSWER_KEY_RE.match(line)
        if mk and cur is not None and not re.search(r"\d+\s*[-–=]", mk.group(1)):
            val = mk.group(1).strip().rstrip(".")
            # strict letter key: single letters only (A / a / А,Б / A,B,C)
            if _LETTER_KEY_RE.match(val) and len(val) <= 12 and cur.answers:
                letters = [c.lower() for c in re.findall(
                    r"[A-Za-zА-Яа-яЇїІіЄєҐґ]", val)]
                _apply_letters(cur, letters)
                continue
            # numeric key: 'Відповідь: 2' marks the 2nd option
            if _NUM_KEY_RE.match(val) and cur.answers:
                nums = [int(x) for x in re.findall(r"\d+", val)]
                for j, a in enumerate(cur.answers):
                    if (j + 1) in nums:
                        a.correct, a.fraction = True, 100
                continue
            # TF answer like "Відповідь: Так"
            if re.fullmatch(r"(?i)(true|false|так|ні|вірно|невірно)", val):
                _make_truefalse(cur, val)
                continue
            # short answer text (question without answer options)
            if cur.answers == [] and val:
                cur.qtype = "shortanswer"
                cur.answers.append(Answer(val, True, 100))
                continue
            continue  # a key line must never merge into the last answer

        # question line? (number must continue the sequence, so numbered
        # answer options like '1) 1945' inside a question are not mistaken)
        qm = _QUESTION_NUM_RE.match(line)
        if qm:
            n = int(qm.group(1))
            has_kw = bool(re.match(
                r"^\s*(?:запитання|питання|question|завдання|q)\b", line,
                re.IGNORECASE))
            qtext = qm.group(2).strip()
            # numbered answer options (1) 2) 3) / 1. 2. 3.) inside a question
            # must NOT look like the next question: they are usually short
            # values (a number, a single word, a year) WITHOUT any '?' or ':'.
            # A numbered line whose text carries '?' or ':' anywhere is a
            # real question, however short it is.
            # A numbered line is an ANSWER OPTION (not a new question) only
            # when it is a numeric-style value (number, year, percentage,
            # measurement) with no '?' and no trailing ':'. A word phrase
            # like '2. При серцево-легеневій реанімації' is a QUESTION even
            # when it is short, because answer variants are lettered
            # (а б в) in the same file; a numbered variant would be numeric.
            wc = len(qtext.split())
            is_numeric_value = bool(
                re.fullmatch(r"[-–\d\.,°%€$'’xх×*=+ ]+", qtext))
            looks_like_option = (
                cur is not None and not has_kw and len(qtext) <= 40
                and not re.search(r"[?]", qtext)
                and not re.search(r":\s*$", qtext)
                and is_numeric_value)
            # a numbered line right after the LAST letter option restarts
            # the variant list (а б в + 1) 2) 3)) → treat as answer option.
            # BUT a line ending in '?' or ':' is a new QUESTION, not an
            # option, even if it follows letter options.
            cur_has_letter_opts = bool(
                cur is not None and cur.answers and all(
                    not getattr(a, "_letter", "").isdigit()
                    for a in cur.answers))
            # A numbered line whose text is question-like (contains '?' or
            # ':' anywhere, OR is a long sentence) is a NEW question, even
            # when it directly follows letter options. Only a short value
            # (number, 1-3 words, no '?'/':') after letter options can be
            # a numeric answer variant like '1) 90°'.
            has_q_mark = bool(re.search(r"[?]", qtext)) or \
                         bool(re.search(r":\s*$", qtext)) or \
                         bool(re.search(r":\s+", qtext))
            # after letter options (а б в), a numbered line is an option
            # only if it is a numeric-style value; a worded sentence is a
            # new question (possibly multi-line).
            is_option_after_letters = (
                cur_has_letter_opts and not has_kw and n <= 20
                and len(qtext) <= 40 and not has_q_mark
                and is_numeric_value)
            cur_has_numeric_opts = bool(
                cur is not None and cur.answers and all(
                    getattr(a, "_letter", "").isdigit()
                    for a in cur.answers))
            is_seq = (not is_option_after_letters) and (
                      has_kw or
                      ((n == last_qnum + 1) and not looks_like_option) or
                      (cur is not None and cur.answers and n > last_qnum
                       and not cur_has_numeric_opts
                       and not looks_like_option) or
                      (cur is None and last_qnum == 0 and n <= 1))
            if is_seq:
                # an unnumbered line right before question 1 with no answers
                # is a document title — drop it; a real unnumbered first
                # question (has answers) is kept by flush()
                if cur is not None and not cur.answers:
                    pending_title = cur.text
                    cur = None
                flush()
                cur = Question(name="", text=qm.group(2).strip())
                cur._qnum = n
                last_qnum = n
                continue

        # '+ а) answer' / '- б) answer' marker style (used in many VNZ keys)
        pm = _MARKER_PREFIX_RE.match(line)
        if pm and cur is not None and cur.text:
            sign, letter, atext = pm.group(1), pm.group(2), pm.group(3)
            atext, marked = _clean_correct_mark(atext)
            is_c = (sign == "+") or marked
            ans = Answer(atext, is_c, 100 if is_c else 0)
            ans._letter = letter.lower()
            cur.answers.append(ans)
            continue

        # unnumbered question: a '?'-terminated line that appears AFTER a
        # completed question (which already has options) starts a new one.
        # This is the layout of medical dumps without any numbering:
        #   Текст питання ... ?
        #   А. варіант
        #   *В. варіант (правильна — зірочка приклеєна до літери)
        if (cur is not None and cur.answers and not questname_mode
                and not _ANSWER_LETTER_RE.match(line)):
            wc = len(line.split())
            # '?'-sentence after options = a new unnumbered question (short
            # questions like 'Хто автор "Гамлета"?' have only 4 words)
            starts_new = bool(_PLAIN_QUESTION_RE.search(line)) and wc >= 3
            # a long sentence after a COMPLETE option block (>=4 variants)
            # is a new question even without '?' (task-style wording that
            # ends with a period: "...Визначте ступінь опіку.")
            if not starts_new and len(cur.answers) >= 4 and wc >= 6:
                starts_new = True
            if starts_new:
                flush()
                cur = Question(name="", text=line.strip())
                cur._qnum = len(questions) + 1
                last_qnum = cur._qnum
                continue

        # answer-variant line?
        am = _ANSWER_LETTER_RE.match(line)
        if am and cur is not None and cur.text:
            star, letter, atext = am.group(1), am.group(2), am.group(3)
            atext, marked = _clean_correct_mark(atext)
            marked = marked or bool(star)   # '*В.' — star glued to letter
            styled = bool(marks & {"bold", "hl", "und", "color"})
            l, r = (None, None)
            if "->" in atext or "=>" in atext or "→" in atext:
                l, r = _split_match(atext)
            ans = Answer(atext, marked or styled,
                         100 if (marked or styled) else 0)
            if l is not None:
                cur.qtype = "matching"
                ans.text, ans.match, ans.correct, ans.fraction = l, r, True, 0
            ans._letter = letter.lower()
            cur.answers.append(ans)
            continue

        # plain continuation of question text (multi-line questions are
        # common in university files: the question spans 2-3 lines).
        # A long sentence line right after a question that has NO options yet
        # is also a continuation of that question's text.
        if cur is not None and not cur.answers and not line.startswith("//"):
            # unnumbered document title: short line without '?' right before
            # the first '?'-question (e.g. 'З номерованими списками') — drop it
            if (not questions and getattr(cur, "_qnum", 1) == 1
                    and len(cur.text) <= 60
                    and not _PLAIN_QUESTION_RE.search(cur.text)
                    and _PLAIN_QUESTION_RE.search(line)):
                pending_title = cur.text
                cur = Question(text=line.strip())
                cur._qnum = 1
                continue
            # PLAIN (unlettered) answer option: the question line is already
            # complete (contains '?') and this bare line carries no '?',
            # e.g. Word auto-numbered/bullet lists whose numbers live in
            # formatting, not text, or plain dumps:
            #   Яка структура координує роботу серця?
            #   головний мозок
            #   *синусовий вузол        (star / bold = correct)
            if (_PLAIN_QUESTION_RE.search(cur.text)
                    and not _PLAIN_QUESTION_RE.search(line)):
                atext, marked = _clean_correct_mark(line.strip())
                styled = bool(marks & {"bold", "hl", "und", "color"})
                cur.answers.append(
                    Answer(atext, marked or styled,
                           100 if (marked or styled) else 0))
            else:
                cur.text += " " + line.strip()
        elif cur is None:
            # file may start with an unnumbered line: either a real first
            # question or just a document title (decided when we see what
            # follows). last_qnum stays 0 so a following '1.' is accepted.
            cur = Question(text=line.strip())
            cur._qnum = len(questions) + 1
        elif cur.answers:
            # continuation of last answer — but NOT if the line actually
            # starts a new unnumbered question. A line containing '?' after
            # a completed option block is ALWAYS a new question (short
            # questions like 'Хто автор "Гамлета"?' have fewer than 6 words).
            wc = len(line.split())
            is_new_q = (not questname_mode
                        and not _ANSWER_LETTER_RE.match(line)
                        and ((_PLAIN_QUESTION_RE.search(line) and wc >= 3)
                             or (len(cur.answers) >= 4 and wc >= 6)))
            if not is_new_q:
                # plain-style option list (unlettered answers): every new
                # bare line is a SEPARATE answer option, not a continuation
                if all(not hasattr(a, "_letter") for a in cur.answers):
                    atext, marked = _clean_correct_mark(line.strip())
                    styled = bool(marks & {"bold", "hl", "und", "color"})
                    cur.answers.append(
                        Answer(atext, marked or styled,
                               100 if (marked or styled) else 0))
                else:
                    cur.answers[-1].text += " " + line.strip()
    flush()

    # fool-protection: a bare list of question lines with NO answers at
    # all is not a valid Moodle test → drop those "questions".
    # Also drop "questions" whose single answer is itself another numbered
    # question line that got merged (sign of a file with no options at all).
    questions = [q for q in questions if q.answers or
                 q.qtype in ("essay", "description")]
    questions = [q for q in questions
                 if not (len(q.answers) == 1 and q.qtype == "multichoice"
                         and re.match(r"^\s*(?:запитання|питання|question|завдання|q)?\s*[№#]?\s*\d{1,4}\s*[\.\)\:\-–—]",
                                      q.answers[0].text, re.IGNORECASE))]

    # apply tail key FIRST — it may be the only source of correct answers
    for q in questions:
        qn = getattr(q, "_qnum", None)
        if qn in tail_key and not any(a.correct for a in q.answers):
            key = tail_key[qn]
            if key.isdigit():
                # numeric key: position of the correct option (1-based)
                idx = int(key) - 1
                if 0 <= idx < len(q.answers):
                    q.answers[idx].correct = True
                    q.answers[idx].fraction = 100
            else:
                for i, a in enumerate(q.answers):
                    if not hasattr(a, "_letter"):
                        a._letter = chr(ord("a") + i)
                _apply_letters(q, [key])
        for a in q.answers:
            if hasattr(a, "_letter"):
                delattr(a, "_letter")

    # fool-protection: if AFTER applying the tail key NO question has a
    # marked correct answer, the file has no usable answer key at all →
    # drop everything rather than silently marking first options correct.
    if questions and not any(a.correct for q in questions for a in q.answers
                             ) and not any(q.qtype in ("essay", "description")
                                           for q in questions):
        questions = []
    # attach the topic title (from 'name:' header) for use as category
    if topic_title:
        for q in questions:
            q._topic = topic_title
    return questions


def _apply_letters(q: Question, letters):
    """Mark answers correct by letter (a/b/c... or а/б/в...)."""
    ua = "абвгдеєжзиіклмнопрстуфхцчшщ"  # ukrainian order
    for i, a in enumerate(q.answers):
        alet = getattr(a, "_letter", None)
        if alet is None:
            alet = chr(ord("a") + i)
        for let in letters:
            hit = (alet == let)
            if not hit and let in ua and alet in ua:
                hit = ua.index(let) == ua.index(alet)
            if not hit and let.isalpha() and alet.isalpha():
                try:
                    hit = (ord(let.lower()) - ord("a")) == \
                          (ord(alet.lower()) - ord("a"))
                except Exception:
                    pass
            if hit:
                a.correct = True
                a.fraction = 100


def _norm_tf(text: str) -> str:
    """Normalize a true/false variant text for comparison."""
    return text.strip().lower().strip("*._ ")


def _make_truefalse(q: Question, val: str):
    v = _norm_tf(val)
    truthy = v in ("true", "так", "вірно", "правда", "1")
    q.qtype = "truefalse"
    q.answers = [Answer("true" if truthy else "false", True, 100)]


def _looks_like_header_row(cells) -> bool:
    joined = " ".join(c.lower() for c in cells)
    return any(k in joined for k in
               ("питання", "question", "відповід", "answer", "correct",
                "правильн", "тип", "type", "№"))


def parse_csv_xlsx_rows(rows):
    """
    Spreadsheet layout. Supported column patterns (auto-detected):
      [question, ans1..ansN, correct(letter or number or text), type?]
    Header row optional; recognised headers: question/питання, answer/відповідь,
    correct/правильна, type/тип.
    """
    if not rows:
        return []
    header = [c.lower() for c in rows[0]]
    has_header = any(k in " ".join(header) for k in
                     ("питання", "question", "відповід", "answer", "correct",
                      "правильн", "тип", "type"))
    data = rows[1:] if has_header else rows

    qi, ti, ci, ni = 0, None, None, None
    if has_header:
        for i, h in enumerate(header):
            if re.fullmatch(r"\s*(№|no|n|номер)\s*\.?\s*", h):
                ni = i
            elif re.search(r"питання|question|запитання", h):
                qi = i
            elif re.search(r"тип|type", h):
                ti = i
            elif re.search(r"correct|правильн|ключ", h):
                ci = i

    questions = []
    for row in data:
        if qi >= len(row) or not row[qi].strip():
            continue
        q = Question(text=row[qi].strip())
        q._qnum = len(questions) + 1
        ans_cols = [i for i in range(len(row))
                    if i not in (qi, ti, ci, ni) and row[i].strip()]
        correct_raw = row[ci].strip() if ci is not None and ci < len(row) else ""
        qtype_raw = row[ti].strip().lower() if ti is not None and ti < len(row) else ""

        if qtype_raw in ("truefalse", "tf", "так/ні"):
            _make_truefalse(q, correct_raw or (row[ans_cols[0]] if ans_cols else ""))
        elif qtype_raw in ("shortanswer", "short", "коротка"):
            q.qtype = "shortanswer"
            vals = [row[i].strip() for i in ans_cols]
            if correct_raw and correct_raw not in vals:
                vals.append(correct_raw)
            q.answers = [Answer(v, True, 100) for v in vals]
        else:
            ua = "абвгдеєжзиіклмнопрстуфхцчшщ"

            def _lidx(ch):
                c = ch.lower()
                if c in ua:
                    return ua.index(c)
                if "a" <= c <= "z":
                    return ord(c) - ord("a")
                return None

            idxs = {i for i in (_lidx(c) for c in
                                re.findall(r"[A-Za-zА-Яа-яЇїІіЄєҐґ]", correct_raw))
                    if i is not None} if correct_raw else set()
            nums = [int(n) for n in re.findall(r"\d+", correct_raw)]
            for j, i in enumerate(ans_cols):
                atext, marked = _clean_correct_mark(row[i].strip())
                is_c = marked or (j in idxs) or ((j + 1) in nums) or \
                       (correct_raw and atext.lower() == correct_raw.lower())
                q.answers.append(Answer(atext, is_c, 100 if is_c else 0))
            # normalize multi-correct weights (header-less sheets etc.)
            corr = [a for a in q.answers if a.correct]
            if len(corr) > 1:
                share = round(100.0 / len(corr), 5)
                for a in corr:
                    a.fraction = share
        questions.append(q)
    return questions


def force_single_answer(questions: List[Question]) -> List[Question]:
    """Force single-correct-answer mode: keep only the FIRST marked answer
    of every multichoice question (100%), the rest become distractors.
    Applied when the user explicitly chooses 'одна правильна відповідь'."""
    for q in questions:
        if q.qtype != "multichoice":
            continue
        correct = [a for a in q.answers if a.correct]
        if len(correct) > 1:
            correct[0].fraction = 100
            for a in correct[1:]:
                a.correct, a.fraction = False, 0
        elif len(correct) == 1:
            correct[0].fraction = 100
    return questions


def post_process(questions: List[Question]) -> List[Question]:
    """Detect question types, fix weights, generate names."""
    for i, q in enumerate(questions, 1):
        if not q.name:
            base = re.sub(r"\s+", " ", q.text).strip()
            q.name = (base[:57] + "...") if len(base) > 60 else base
            q.name = q.name or f"Q{i}"
        if q.qtype == "description":
            continue
        if q.qtype == "matching":
            for a in q.answers:
                a.correct, a.fraction = True, 0
            continue
        # numeric?
        if q.qtype == "multichoice" and q.answers and all(
                a.correct and re.fullmatch(r"-?\d+([.,]\d+)?", a.text.replace(",", "."))
                for a in q.answers if a.correct) and \
                sum(1 for a in q.answers if a.correct) == len(q.answers) == 1:
            q.qtype = "numerical"
            continue
        # true/false by answers content (ua + en, incl. unmarked pairs)
        if q.qtype == "multichoice" and len(q.answers) == 2:
            pair = {_norm_tf(a.text) for a in q.answers}
            tf_pairs = ({"true", "false"}, {"так", "ні"},
                        {"вірно", "невірно"}, {"правда", "брехня"})
            if any(pair == p for p in tf_pairs):
                corr = next((a for a in q.answers if a.correct), None)
                if corr is None:
                    # default: the affirmative variant is the truth value
                    corr = next(a for a in q.answers if _norm_tf(a.text)
                                in ("true", "так", "вірно", "правда"))
                _make_truefalse(q, corr.text)
                continue
        # all answers marked correct: if there are 2+ options this is a
        # multiple-response question (multi-select), NOT a short answer.
        # Short answer = exactly ONE option with no distractors.
        if q.qtype == "multichoice" and q.answers and \
                all(a.correct for a in q.answers):
            if len(q.answers) == 1:
                q.qtype = "shortanswer"
            # else: keep multichoice, fractions are distributed below
            if q.qtype == "shortanswer":
                continue
        # multichoice weights (keep fractions already set, e.g. from sheets)
        correct = [a for a in q.answers if a.correct]
        wrong = [a for a in q.answers if not a.correct]
        if not correct and q.answers:
            q.answers[0].correct = True
            q.answers[0].fraction = 100
            correct = [q.answers[0]]
        q._multi = len(correct) > 1  # remember multi-select flag
        if len(correct) == 1:
            correct[0].fraction = 100
            for a in wrong:
                if a.fraction > 0:
                    a.fraction = 0
        elif correct:
            # multi-select: always (re)distribute 100% evenly across the
            # correct options (weights may already be set to 100 each at
            # parse time — normalize them here)
            share = round(100.0 / len(correct), 5)
            for a in correct:
                a.fraction = share
            for a in wrong:
                a.fraction = 0
    return questions


def parse_file(path: str) -> List[Question]:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".docx":
        paras, tables = read_docx_paragraphs(path)
        lines = [p[0] for p in paras]
        marks = [p[1] for p in paras]
        qs = parse_numbered_lines(lines, marks)
        # Word table templates (№ | Питання | Відповіді... | Правильна)
        if tables:
            tqs = parse_csv_xlsx_rows(tables)
            if tqs:
                qs = qs + tqs if qs else tqs
    elif ext in (".xlsx", ".xlsm"):
        qs = parse_csv_xlsx_rows(read_xlsx_rows(path))
    elif ext == ".csv":
        qs = parse_csv_xlsx_rows(read_csv_rows(path))
    elif ext == ".doc":
        qs = parse_numbered_lines(read_doc_text(path).splitlines())
    elif ext in (".html", ".htm"):
        qs = parse_numbered_lines(read_html_text(path).splitlines())
    else:
        qs = parse_numbered_lines(read_text_file(path).splitlines())
    return post_process(qs)


# ---------------------------------------------------------------------------
# Exporters
# ---------------------------------------------------------------------------

_GIFT_SPECIALS = ("~", "=", "#", "{", "}", ":")


def gift_escape(text: str) -> str:
    out = []
    for ch in text:
        if ch in _GIFT_SPECIALS or ch == "\\":
            out.append("\\" + ch)
        else:
            out.append(ch)
    return "".join(out).replace("\n", "\\n")


def _fmt_frac(f: float) -> str:
    s = f"{f:.5f}".rstrip("0").rstrip(".")
    return s


def to_gift(questions: List[Question], category: str = "") -> str:
    parts = []
    if category:
        parts.append(f"$CATEGORY: {category}\n")
    for q in questions:
        if q.qtype == "description":
            parts.append(f"::{gift_escape(q.name)}::{gift_escape(q.text)}\n")
            continue
        head = f"::{gift_escape(q.name)}::{gift_escape(q.text)}"
        if q.feedback:
            head += f" ####{gift_escape(q.feedback)}"
        if q.qtype == "truefalse":
            val = "T" if q.answers[0].text.lower() in ("true", "так", "вірно") else "F"
            parts.append(f"{head} {{{val}}}\n")
        elif q.qtype == "shortanswer":
            body = " ".join(f"={gift_escape(a.text)}" for a in q.answers)
            parts.append(f"{head} {{{body}}}\n")
        elif q.qtype == "matching":
            body = "\n".join(
                f"  ={gift_escape(a.text)} -> {gift_escape(a.match)}"
                for a in q.answers)
            parts.append(f"{head} {{\n{body}\n}}\n")
        elif q.qtype == "numerical":
            a = q.answers[0]
            parts.append(f"{head} {{#{a.text.replace(',', '.')}}}\n")
        elif q.qtype == "essay":
            parts.append(f"{head} {{}}\n")
        else:  # multichoice
            lines = []
            for a in q.answers:
                fb = f" #{gift_escape(a.feedback)}" if a.feedback else ""
                if a.correct and a.fraction >= 100:
                    lines.append(f"  ={gift_escape(a.text)}{fb}")
                elif a.correct:
                    lines.append(
                        f"  ~%{_fmt_frac(a.fraction)}%{gift_escape(a.text)}{fb}")
                elif a.fraction < 0:
                    lines.append(
                        f"  ~%{_fmt_frac(a.fraction)}%{gift_escape(a.text)}{fb}")
                else:
                    lines.append(f"  ~{gift_escape(a.text)}{fb}")
            parts.append(head + " {\n" + "\n".join(lines) + "\n}\n")
    return "\n".join(parts)


def to_aiken(questions: List[Question]) -> str:
    blocks = []
    for q in questions:
        if q.qtype != "multichoice":
            continue  # Aiken supports only single-answer MC
        correct = [a for a in q.answers if a.correct]
        if len(correct) != 1:
            continue
        lines = [q.text.strip()]
        ci = 0
        for i, a in enumerate(q.answers):
            lines.append(f"{chr(65 + i)}. {a.text}")
            if a.correct:
                ci = i
        lines.append(f"ANSWER: {chr(65 + ci)}")
        blocks.append("\n".join(lines))
    return "\n\n".join(blocks) + "\n"


def to_moodle_xml(questions: List[Question], category: str = "") -> str:
    x = ['<?xml version="1.0" encoding="UTF-8"?>', "<quiz>"]
    if category:
        x.append(f'  <question type="category"><category>'
                 f'<text>{saxutils.escape(category)}</text></category>'
                 f"</question>")
    for q in questions:
        name = saxutils.escape(q.name)
        text = saxutils.escape(q.text)
        gf = saxutils.escape(q.feedback)
        mark = q.default_mark
        if q.qtype == "description":
            x.append(f'''  <question type="description">
    <name><text>{name}</text></name>
    <questiontext format="moodle_auto_format"><text>{text}</text></questiontext>
  </question>''')
        elif q.qtype == "truefalse":
            val = q.answers[0].text.lower() in ("true", "так", "вірно")
            x.append(f'''  <question type="truefalse">
    <name><text>{name}</text></name>
    <questiontext format="moodle_auto_format"><text>{text}</text></questiontext>
    <generalfeedback format="moodle_auto_format"><text>{gf}</text></generalfeedback>
    <defaultgrade>{mark}</defaultgrade>
    <answer fraction="{'100' if val else '0'}" format="moodle_auto_format"><text>true</text></answer>
    <answer fraction="{'0' if val else '100'}" format="moodle_auto_format"><text>false</text></answer>
  </question>''')
        elif q.qtype == "shortanswer":
            ans = "\n".join(
                f'    <answer fraction="100" format="moodle_auto_format">'
                f"<text>{saxutils.escape(a.text)}</text></answer>"
                for a in q.answers)
            x.append(f'''  <question type="shortanswer">
    <name><text>{name}</text></name>
    <questiontext format="moodle_auto_format"><text>{text}</text></questiontext>
    <generalfeedback format="moodle_auto_format"><text>{gf}</text></generalfeedback>
    <defaultgrade>{mark}</defaultgrade>
    <usecase>0</usecase>
{ans}
  </question>''')
        elif q.qtype == "matching":
            subs = "\n".join(
                f"    <subquestion format=\"moodle_auto_format\">"
                f"<text>{saxutils.escape(a.text)}</text>"
                f"<answer><text>{saxutils.escape(a.match)}</text></answer>"
                f"</subquestion>" for a in q.answers)
            x.append(f'''  <question type="matching">
    <name><text>{name}</text></name>
    <questiontext format="moodle_auto_format"><text>{text}</text></questiontext>
    <defaultgrade>{mark}</defaultgrade>
    <shuffleanswers>true</shuffleanswers>
{subs}
  </question>''')
        elif q.qtype == "numerical":
            a = q.answers[0]
            x.append(f'''  <question type="numerical">
    <name><text>{name}</text></name>
    <questiontext format="moodle_auto_format"><text>{text}</text></questiontext>
    <defaultgrade>{mark}</defaultgrade>
    <answer fraction="100" format="moodle_auto_format">
      <text>{saxutils.escape(a.text.replace(',', '.'))}</text>
      <tolerance>0</tolerance>
    </answer>
  </question>''')
        elif q.qtype == "essay":
            x.append(f'''  <question type="essay">
    <name><text>{name}</text></name>
    <questiontext format="moodle_auto_format"><text>{text}</text></questiontext>
    <defaultgrade>{mark}</defaultgrade>
    <responseformat>editor</responseformat>
  </question>''')
        else:  # multichoice
            correct_n = sum(1 for a in q.answers if a.correct)
            single = "true" if correct_n <= 1 else "false"
            ans = []
            for a in q.answers:
                fb = (f'<feedback format="moodle_auto_format"><text>'
                      f"{saxutils.escape(a.feedback)}</text></feedback>"
                      if a.feedback else "")
                ans.append(
                    f'    <answer fraction="{_fmt_frac(a.fraction)}" '
                    f'format="moodle_auto_format"><text>'
                    f"{saxutils.escape(a.text)}</text>{fb}</answer>")
            x.append(f'''  <question type="multichoice">
    <name><text>{name}</text></name>
    <questiontext format="moodle_auto_format"><text>{text}</text></questiontext>
    <generalfeedback format="moodle_auto_format"><text>{gf}</text></generalfeedback>
    <defaultgrade>{mark}</defaultgrade>
    <single>{single}</single>
    <shuffleanswers>true</shuffleanswers>
    <answernumbering>abc</answernumbering>
{chr(10).join(ans)}
  </question>''')
    x.append("</quiz>")
    return "\n".join(x) + "\n"


def convert(path: str, out_format: str, category: str = "",
            single: bool = False) -> str:
    qs = parse_file(path)
    if single:
        force_single_answer(qs)
    # 'name:' header of a newTem; dump overrides the default category
    if not category and qs:
        category = getattr(qs[0], "_topic", "") or ""
    if out_format == "gift":
        return to_gift(qs, category)
    if out_format == "aiken":
        return to_aiken(qs)
    return to_moodle_xml(qs, category)
